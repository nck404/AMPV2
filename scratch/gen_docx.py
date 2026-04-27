from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_word_file():
    doc = Document()
    
    # Title
    title = doc.add_heading('MÔ TẢ CÁC TÍNH NĂNG CƠ BẢN CỦA SẢN PHẨM AMP V2', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Content from txt file
    content = """
1. Mô tả các tính năng cơ bản của sản phẩm
Sản phẩm AMP là một nền tảng đa mục tiêu, hướng tới việc hỗ trợ toàn diện cho cộng đồng, đặc biệt là người khuyết tật và người lao động phổ thông. Các chức năng chính bao gồm:

* Hệ thống Tuyển dụng & Tạo hồ sơ số (CV): Cho phép người dùng tìm kiếm việc làm theo ngành nghề, địa điểm và hình thức làm việc. Tính năng tạo CV tự động giúp người dùng có ngay hồ sơ chuyên nghiệp chỉ bằng cách điền thông tin cơ bản.
* Học tập Ngôn ngữ ký hiệu tích hợp AI: Cung cấp lộ trình học ngôn ngữ ký hiệu (ASL/VDSL) từ cơ bản đến nâng cao. Điểm đặc biệt là hệ thống sử dụng AI để nhận diện cử chỉ tay qua camera thời gian thực, giúp người dùng thực hành và nhận phản hồi ngay lập tức.
* Bộ công cụ Trợ năng (Accessibility Tools): 
    * Text-to-Speech (TTS): Chuyển đổi văn bản thành giọng nói để hỗ trợ người khiếm thị.
    * Nhận diện mệnh giá tiền/số: Hỗ trợ người khiếm thị và người già trong giao dịch hàng ngày.
    * Giao diện đọc sách đa phương tiện: Hỗ trợ đọc tài liệu kết hợp âm thanh (Audiobook) và video minh họa.
* Diễn đàn Cộng đồng & Kết nối: Không gian thảo luận văn minh, nơi người dùng chia sẻ kinh nghiệm nghề nghiệp, kết bạn và nhắn tin trực tiếp để mở ra các cơ hội hợp tác.
* Hệ thống Quản trị (Admin Dashboard): Công cụ quản lý nội dung mạnh mẽ cho phép quản trị viên cập nhật tin tuyển dụng, bài học và tài liệu minh họa một cách dễ dàng.

2. Ưu điểm và Tính mới của sản phẩm
* Tính toàn diện (All-in-one Ecosystem): Khác với các sản phẩm giáo dục hoặc tuyển dụng đơn thuần, AMP kết hợp cả ba yếu tố: Công cụ trợ năng - Giáo dục - Việc làm vào một hệ sinh thái duy nhất, tạo ra một lộ trình hòa nhập trọn vẹn cho người yếu thế.
* Ứng dụng AI đột phá: Sử dụng AI trực tiếp trên trình duyệt (Edge Computing) để nhận diện ngôn ngữ ký hiệu. Điều này giúp tối ưu tốc độ phản hồi (dưới 0.8s) và đảm bảo tính riêng tư vì dữ liệu camera không cần gửi về server.
* Thiết kế Accessibility-First: Giao diện được thiết kế theo phong cách hiện đại (Premium Design), nhưng tuân thủ nghiêm ngặt các tiêu chuẩn về độ tương phản, kích thước chữ và điều hướng bằng bàn phím để người khiếm thị và người hạn chế vận động có thể sử dụng dễ dàng.
* Ý tưởng mới về "Hòa nhập số": Sản phẩm không chỉ hỗ trợ người khuyết tật dùng công nghệ, mà còn dùng công nghệ để giúp họ tạo ra giá trị kinh tế thông qua kênh tuyển dụng riêng biệt.

3. Thành phần sử dụng lại và Phát triển mới
Sản phẩm được xây dựng dựa trên sự kết hợp giữa các công nghệ mã nguồn mở và phần tự phát triển mẫu chốt:
* Mã nguồn sử dụng lại (Open Source):
    * Frontend Framework: Thư viện SvelteKit & TailwindCSS (Tạo cấu trúc và giao diện).
    * AI Hand Tracking: Sử dụng thư viện MediaPipe Hands của Google và Fingerpose (Dùng cho chức năng nhận diện cử chỉ tay). Link: https://github.com/google/mediapipe.
    * Backend: Flask (Python) kèm các thư viện hỗ trợ như gTTS (Text-to-Speech).
* Chức năng tự cài đặt (Custom Development):
    * Logic nhận diện ngôn ngữ ký hiệu: Tự xây dựng bộ quy tắc (Classifiers) cho các chữ cái và từ vựng dựa trên tọa độ xương tay từ MediaPipe.
    * Hệ thống quản lý Multimedia: Toàn bộ backend xử lý upload, chuyển đổi định dạng sách và video được lập trình riêng.
    * Công cụ tạo CV số: Hệ thống tự động chuyển đổi dữ liệu người dùng thành các mẫu CV chuyên nghiệp.
"""
    
    for line in content.split('\n'):
        if not line.strip():
            continue
        
        if line.startswith('1.') or line.startswith('2.') or line.startswith('3.'):
            heading = doc.add_heading(line, level=1)
        elif line.startswith('*'):
            doc.add_paragraph(line.replace('*', '').strip(), style='List Bullet')
        elif line.startswith('    *'):
            p = doc.add_paragraph(line.replace('*', '').strip(), style='List Bullet 2')
        else:
            doc.add_paragraph(line)
            
    doc.save('d:/Project/AMPV2/mo_ta_san_pham_AMPV2.docx')

if __name__ == "__main__":
    create_word_file()
